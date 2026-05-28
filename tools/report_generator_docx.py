from __future__ import annotations

import argparse
import json
import re
from collections import Counter, defaultdict
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = ROOT / "templates/initial_report_en.docx"
DEFAULT_DATA = ROOT / "generated_reports/report_data.json"
DEFAULT_OUTPUT = ROOT / "generated_reports/initial_report_en.docx"

SEVERITIES = ["Critical", "High", "Medium", "Low", "Informational"]


def normalize_key(value: str) -> str:
    return re.sub(r"[^A-Z0-9]+", "_", value.upper()).strip("_")


def iter_block_items(doc: Document):
    parent = doc.element.body
    for child in parent.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def element_text(element) -> str:
    text_nodes = [node.text or "" for node in element.iter(qn("w:t"))]
    if text_nodes:
        return "".join(text_nodes).strip()
    return "".join(element.itertext()).strip()


def body_children(doc: Document) -> list[Any]:
    return [child for child in doc.element.body.iterchildren() if child.tag != qn("w:sectPr")]


def find_marker_index(children: list[Any], marker: str, start: int = 0) -> int:
    for idx in range(start, len(children)):
        if is_marker_element(children[idx], marker):
            return idx
    return -1


def is_marker_element(element, marker: str) -> bool:
    text = element_text(element)
    if text == marker:
        return True
    markers = re.findall(r"\{\{[^{}]+\}\}", text)
    return bool(markers) and all(item == marker for item in markers)


def delete_range(body, start: int, end: int) -> None:
    for child in list(body)[start : end + 1]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def make_page_break_paragraph():
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def has_page_break(element) -> bool:
    xml = etree_tostring(element)
    return "<w:br" in xml and 'w:type="page"' in xml


def is_page_break_only(element) -> bool:
    if element.tag != qn("w:p"):
        return False
    text = re.sub(r"\s+", "", element_text(element))
    return not text and has_page_break(element)


def etree_tostring(element) -> str:
    from lxml import etree

    return etree.tostring(element, encoding="unicode")


def normalize_findings_table_breaks(doc: Document) -> None:
    """Keep Findings Table continuous, then force Details of Findings to a new page."""
    body = doc.element.body
    children = body_children(doc)

    # Only match actual paragraphs, not structured document tags (TOC contains heading text too)
    def is_heading_paragraph(child, heading_text):
        return child.tag == qn("w:p") and heading_text in element_text(child)

    findings_idx = next((i for i, child in enumerate(children) if is_heading_paragraph(child, "FINDINGS TABLE")), -1)
    details_idx = next(
        (i for i, child in enumerate(children[findings_idx + 1 :], start=findings_idx + 1) if is_heading_paragraph(child, "DETAILS OF FINDINGS")),
        -1,
    )
    if findings_idx < 0 or details_idx < 0 or details_idx <= findings_idx:
        return

    for child in list(children[findings_idx + 1 : details_idx]):
        if is_page_break_only(child):
            body.remove(child)

    children = body_children(doc)
    findings_idx = next((i for i, child in enumerate(children) if is_heading_paragraph(child, "FINDINGS TABLE")), -1)
    details_idx = next(
        (i for i, child in enumerate(children[findings_idx + 1 :], start=findings_idx + 1) if is_heading_paragraph(child, "DETAILS OF FINDINGS")),
        -1,
    )
    if details_idx > 0 and not is_page_break_only(children[details_idx - 1]):
        body.insert(details_idx, make_page_break_paragraph())


def set_keep_with_next(element) -> None:
    if element.tag != qn("w:p"):
        return
    p_pr = element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        element.insert(0, p_pr)
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def keep_finding_headings_together(doc: Document) -> None:
    """Avoid orphan severity headings when Word/Google Docs repaginates."""
    severity_headings = tuple(f"{severity}-Risk Findings" for severity in SEVERITIES)
    children = body_children(doc)
    for idx, child in enumerate(children):
        text = re.sub(r"\s+", " ", element_text(child)).strip()
        if any(heading in text for heading in severity_headings):
            set_keep_with_next(child)
            next_idx = idx + 1
            while next_idx < len(children) and not element_text(children[next_idx]).strip():
                next_idx += 1
            if next_idx < len(children) and children[next_idx].tag == qn("w:p"):
                set_keep_with_next(children[next_idx])


def force_major_section_page_breaks(doc: Document) -> None:
    section_headings = (
        "CONFIDENTIALITY",
        "THANK YOU",
        "TABLE OF CONTENTS",
        "DOCUMENT DETAILS",
        "SECURITY ASSESSMENT INFORMATION",
        "PROJECT DETAILS",
        "Methodology",
        "Classification of Level of Vulnerability",
        "EXECUTIVE SUMMARY",
        "FINDINGS TABLE",
        # "DETAILS OF FINDINGS" is handled by normalize_findings_table_breaks()
        "Appendix A: Scope Details",
    )
    body = doc.element.body
    idx = 1
    while idx < len(body_children(doc)):
        children = body_children(doc)
        child = children[idx]
        if child.tag != qn("w:p"):
            idx += 1
            continue
        text = re.sub(r"\s+", " ", element_text(child)).strip()
        compact_text = re.sub(r"[^A-Za-z0-9]+", "", text).upper()
        is_heading = any(
            compact_text in {re.sub(r"[^A-Za-z0-9]+", "", heading).upper() * repeat for repeat in range(1, 5)}
            for heading in section_headings
        )
        if not is_heading:
            idx += 1
            continue
        prev = children[idx - 1]
        if not is_page_break_only(prev) and not has_page_break(prev):
            body.insert(idx, make_page_break_paragraph())
            idx += 2
        else:
            idx += 1


def replace_text_in_element(element, context: dict[str, Any]) -> None:
    # Track placeholders already replaced in this element. Google Docs / Word can
    # duplicate the same placeholder inside one w:t node or across several runs.
    # Only the first occurrence should receive the value; later occurrences are
    # removed so generated prose is not repeated two or three times.
    replaced_keys: set[str] = set()
    for text_node in element.iter(qn("w:t")):
        if not text_node.text:
            continue
        value = text_node.text
        for key, replacement in context.items():
            marker = "{{" + key + "}}"
            if marker not in value:
                continue
            if key in replaced_keys:
                value = value.replace(marker, "")
            else:
                value = value.replace(marker, str(replacement), 1)
                value = value.replace(marker, "")
                replaced_keys.add(key)
        text_node.text = value


def exact_repeated_unit(text: str) -> str | None:
    compact = text.strip()
    if len(compact) < 12:
        return None
    for repeat_count in (3, 2):
        if len(compact) % repeat_count:
            continue
        unit = compact[: len(compact) // repeat_count]
        if unit and unit * repeat_count == compact:
            return unit
    return None


def set_text_in_element(element, text: str) -> None:
    first = True
    for text_node in element.iter(qn("w:t")):
        if first:
            text_node.text = text
            text_node.set(qn("xml:space"), "preserve")
            first = False
        else:
            text_node.text = ""


def collapse_exact_repeated_paragraphs(doc: Document) -> None:
    """Clean Google Docs import artifacts like 'texttexttext' in one paragraph."""
    for child in body_children(doc):
        if child.tag != qn("w:p"):
            continue
        text = element_text(child)
        if "{{" in text or text.startswith("__IMAGE__:"):
            continue
        unit = exact_repeated_unit(text)
        if unit:
            set_text_in_element(child, unit)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def insert_image_placeholders(doc: Document) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text.startswith("__IMAGE__:"):
            # Center figure captions (e.g., "Figure 1. ...")
            if re.match(r"Figure\s+\d+", text):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        path = Path(text.removeprefix("__IMAGE__:"))
        clear_paragraph(paragraph)
        run = paragraph.add_run()
        if path.exists():
            run.add_picture(str(path), width=docx_inches(5.9))
        else:
            run.add_text(f"[missing image: {path}]")
        # Center-align the image paragraph
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _split_text_into_lines(text: str) -> list[str]:
    """Split a text block into lines, handling:
    - Actual newlines
    - Numbered lists: '.1. ', '.2. ' etc (period before number)
    - Bullet lists: '.- ', ':- ' (after period or colon)
    """
    if not text:
        return [text]

    # First, split on actual newlines
    if "\n" in text:
        parts = [line.strip() for line in text.split("\n") if line.strip()]
        if len(parts) > 1:
            return parts

    # Split numbered lists: e.g., "...command.2. Utilize..." -> ["...command.", "2. Utilize..."]
    # Pattern: period/colon followed by a number and period
    split = re.split(r"(?<=[\.\:\;])\s*(?=\d+\.\s)", text)
    if len(split) > 1:
        return [s.strip() for s in split if s.strip()]

    # Split bullet lists: e.g., "...following:- Item one.- Item two"
    split = re.split(r"(?<=[\.\:\;])\s*(?=-\s)", text)
    if len(split) > 1:
        return [s.strip() for s in split if s.strip()]

    return [text]


def _clone_paragraph_with_new_text(original_element, new_text: str):
    """Create a new paragraph element with the same formatting but different text."""
    new_p = deepcopy(original_element)
    # Clear all runs' text and set the first run only, otherwise Google Docs
    # imported templates can keep duplicate text nodes alive after splitting.
    first_t = next(new_p.iter(qn("w:t")), None)
    if first_t is not None:
        set_text_in_element(new_p, new_text)
    else:
        # No text node found, create a run with text
        run_el = OxmlElement("w:r")
        # Copy run properties from original if available
        orig_run = next(original_element.iter(qn("w:r")), None)
        if orig_run is not None:
            rpr = orig_run.find(qn("w:rPr"))
            if rpr is not None:
                run_el.append(deepcopy(rpr))
        t_el = OxmlElement("w:t")
        t_el.text = new_text
        t_el.set(qn("xml:space"), "preserve")
        run_el.append(t_el)
        new_p.append(run_el)
    return new_p


def split_multiline_paragraphs(doc: Document) -> None:
    """Split paragraphs containing multi-line text (newlines, numbered/bullet lists)
    into separate Word paragraphs, preserving formatting."""
    body = doc.element.body
    children = body_children(doc)
    # Process in reverse so insertions don't shift indices
    for idx in range(len(children) - 1, -1, -1):
        child = children[idx]
        if child.tag != qn("w:p"):
            continue
        text = element_text(child)
        if not text or len(text) < 20:
            continue
        # Skip special markers/headings
        if text.startswith("__IMAGE__:") or text.startswith("{{"):
            continue

        lines = _split_text_into_lines(text)
        if len(lines) <= 1:
            continue

        # Replace original paragraph with multiple paragraphs
        parent = child.getparent()
        insert_point = list(parent).index(child)
        parent.remove(child)
        for offset, line in enumerate(lines):
            new_p = _clone_paragraph_with_new_text(child, line)
            parent.insert(insert_point + offset, new_p)


def docx_inches(value: float):
    from docx.shared import Inches

    return Inches(value)


def replace_paragraph_with_image_sentinel(element, path: str) -> None:
    resolved_path = Path(path)
    if not resolved_path.is_absolute():
        resolved_path = ROOT / str(path).lstrip("/")
    replaced = False
    for text_node in element.iter(qn("w:t")):
        if not text_node.text or "{{POC_IMAGE}}" not in text_node.text:
            continue
        if replaced:
            text_node.text = text_node.text.replace("{{POC_IMAGE}}", "")
            continue
        text_node.text = text_node.text.replace("{{POC_IMAGE}}", "__IMAGE__:" + str(resolved_path), 1)
        text_node.text = text_node.text.replace("{{POC_IMAGE}}", "")
        replaced = True


def build_context(data: dict[str, Any]) -> dict[str, Any]:
    findings = data.get("findings", [])
    counts = Counter(f.get("severity", "") for f in findings)
    context = {normalize_key(key): value for key, value in data.items() if not isinstance(value, (list, dict))}
    context.update(
        {
            "TARGET_TOTAL": len(data.get("scopes", [])),
            "CRITICAL_COUNT": counts.get("Critical", 0),
            "HIGH_COUNT": counts.get("High", 0),
            "MEDIUM_COUNT": counts.get("Medium", 0),
            "LOW_COUNT": counts.get("Low", 0),
            "INFORMATIONAL_COUNT": counts.get("Informational", 0),
            # Fallbacks for stale Table of Contents snapshots. The real TOC can be refreshed
            # in Word/Google Docs after generation.
            "FINDING_NAME": "",
        }
    )
    return context


def finding_context(finding: dict[str, Any]) -> dict[str, Any]:
    return {normalize_key(key): value for key, value in finding.items() if not isinstance(value, (list, dict))}


def process_condition_blocks(doc: Document, findings_by_severity: dict[str, list[dict[str, Any]]]) -> None:
    body = doc.element.body
    for severity in SEVERITIES:
        block = f"HAS_{normalize_key(severity)}_FINDINGS"
        has_data = bool(findings_by_severity.get(severity))
        while True:
            children = body_children(doc)
            open_idx = find_marker_index(children, "{{#" + block + "}}")
            if open_idx < 0:
                break
            close_idx = find_marker_index(children, "{{/" + block + "}}", open_idx + 1)
            if close_idx < 0:
                raise ValueError(f"Missing closing marker for {block}")
            if has_data:
                body.remove(children[close_idx])
                body.remove(children[open_idx])
            else:
                delete_range(body, open_idx, close_idx)

        while True:
            children = body_children(doc)
            inv_idx = find_marker_index(children, "{{^" + block + "}}")
            if inv_idx < 0:
                break
            close_idx = find_marker_index(children, "{{/" + block + "}}", inv_idx + 1)
            if close_idx < 0:
                raise ValueError(f"Missing inverse closing marker for {block}")
            if has_data:
                delete_range(body, inv_idx, close_idx)
            else:
                body.remove(children[close_idx])
                body.remove(children[inv_idx])


def process_body_repeat_block(
    doc: Document,
    block: str,
    rows: list[dict[str, Any]],
    mapper,
    *,
    add_page_break_between_items: bool = False,
    page_break_state: dict[str, int] | None = None,
) -> None:
    body = doc.element.body
    while True:
        children = body_children(doc)
        open_idx = find_marker_index(children, "{{#" + block + "}}")
        if open_idx < 0:
            return
        close_idx = find_marker_index(children, "{{/" + block + "}}", open_idx + 1)
        if close_idx < 0:
            raise ValueError(f"Missing closing marker for {block}")
        template = [deepcopy(child) for child in children[open_idx + 1 : close_idx]]
        output = []
        for row in rows:
            context = mapper(row)
            clones = [deepcopy(child) for child in template]
            process_nested_repeat_blocks(clones, row)
            if add_page_break_between_items:
                if page_break_state is None:
                    page_break_state = {"count": 0}
                if page_break_state["count"] > 0:
                    output.append(make_page_break_paragraph())
                page_break_state["count"] += 1
            for clone in clones:
                replace_text_in_element(clone, context)
                output.append(clone)
        for child in children[open_idx : close_idx + 1]:
            body.remove(child)
        for offset, child in enumerate(output):
            body.insert(open_idx + offset, child)


def process_nested_repeat_blocks(elements: list[Any], finding: dict[str, Any]) -> None:
    repeat_specs = {
        "AFFECTED_ITEMS": [{"AFFECTED_ITEM": item} for item in finding.get("affected_items", [])],
        "REFERENCES": [{"REFERENCE": item} for item in finding.get("references", [])],
        "POC_IMAGES": [
            {
                "POC_IMAGE": image.get("image_path", ""),
                "FIGURE_NUMBER": idx,
                "FIGURE_CAPTION": image.get("figure_caption", ""),
            }
            for idx, image in enumerate(finding.get("poc_images", []), start=1)
        ],
    }
    for block, rows in repeat_specs.items():
        while True:
            open_idx = next((i for i, el in enumerate(elements) if is_marker_element(el, "{{#" + block + "}}")), -1)
            if open_idx < 0:
                break
            close_idx = next((i for i, el in enumerate(elements[open_idx + 1 :], start=open_idx + 1) if is_marker_element(el, "{{/" + block + "}}")), -1)
            if close_idx < 0:
                raise ValueError(f"Missing nested closing marker for {block}")
            template = [deepcopy(el) for el in elements[open_idx + 1 : close_idx]]
            output = []
            for row in rows:
                for clone in [deepcopy(el) for el in template]:
                    if block == "POC_IMAGES":
                        replace_paragraph_with_image_sentinel(clone, row["POC_IMAGE"])
                    replace_text_in_element(clone, row)
                    output.append(clone)
            elements[open_idx : close_idx + 1] = output


def process_table_repeat_blocks(doc: Document, data: dict[str, Any], findings_by_severity: dict[str, list[dict[str, Any]]]) -> None:
    table_specs: dict[str, list[dict[str, Any]]] = {
        "SCOPES": [
            {
                "SCOPE_NAME": scope.get("scope_name", ""),
                "SCOPE_TARGET": scope.get("scope_target", ""),
                "SCOPE_AREA": scope.get("scope_area", ""),
                "TESTING_APPROACH": data.get("testing_approach", ""),
            }
            for scope in data.get("scopes", [])
        ]
    }
    for severity in SEVERITIES:
        block = f"{normalize_key(severity)}_FINDINGS_TABLE"
        table_specs[block] = [finding_context(finding) for finding in findings_by_severity.get(severity, [])]

    for table in doc.tables:
        tbl = table._tbl
        for block, rows in table_specs.items():
            while True:
                tr_list = list(tbl.tr_lst)
                open_idx = next((i for i, row in enumerate(tr_list) if "{{#" + block + "}}" in element_text(row)), -1)
                if open_idx < 0:
                    break
                close_idx = next((i for i, row in enumerate(tr_list[open_idx + 1 :], start=open_idx + 1) if "{{/" + block + "}}" in element_text(row)), -1)
                if close_idx < 0:
                    raise ValueError(f"Missing table closing marker for {block}")
                template_rows = [deepcopy(row) for row in tr_list[open_idx + 1 : close_idx]]
                output_rows = []
                for row_data in rows:
                    for clone in [deepcopy(row) for row in template_rows]:
                        replace_text_in_element(clone, row_data)
                        output_rows.append(clone)
                for row in tr_list[open_idx : close_idx + 1]:
                    tbl.remove(row)
                if open_idx > 0:
                    anchor = list(tbl.tr_lst)[open_idx - 1]
                    for row in output_rows:
                        anchor.addnext(row)
                        anchor = row
                else:
                    for row in reversed(output_rows):
                        tbl.insert(0, row)


def replace_global_placeholders(doc: Document, context: dict[str, Any]) -> None:
    for child in doc.element.body.iterchildren():
        replace_text_in_element(child, context)


def unresolved_markers(doc: Document) -> list[str]:
    text = "\n".join(element_text(child) for child in doc.element.body.iterchildren())
    return sorted(set(re.findall(r"\{\{[^{}]+\}\}", text)))


def generate(template: Path, data_path: Path, output: Path) -> dict[str, Any]:
    data = json.loads(data_path.read_text())
    doc = Document(template)
    findings_by_severity: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for finding in data.get("findings", []):
        findings_by_severity[finding.get("severity", "")].append(finding)

    process_condition_blocks(doc, findings_by_severity)
    # Per-severity page_break_state so the severity heading flows naturally
    # into its first finding (no orphan heading pages)
    for severity in SEVERITIES:
        block = f"{normalize_key(severity)}_FINDINGS"
        process_body_repeat_block(
            doc,
            block,
            findings_by_severity.get(severity, []),
            finding_context,
            add_page_break_between_items=True,
            page_break_state={"count": 0},
        )
    process_table_repeat_blocks(doc, data, findings_by_severity)
    normalize_findings_table_breaks(doc)
    keep_finding_headings_together(doc)
    force_major_section_page_breaks(doc)
    replace_global_placeholders(doc, build_context(data))
    collapse_exact_repeated_paragraphs(doc)
    # Only split multi-line paragraphs for English reports
    if "_en" in template.stem:
        split_multiline_paragraphs(doc)
        collapse_exact_repeated_paragraphs(doc)
    insert_image_placeholders(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    unresolved = unresolved_markers(doc)
    return {
        "output": str(output),
        "findings": len(data.get("findings", [])),
        "unresolved": unresolved,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate an initial pentest report DOCX from VulnVault project data.")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--data", type=Path, default=DEFAULT_DATA)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    result = generate(args.template, args.data, args.out)
    print(json.dumps(result, indent=2))
    return 0 if not result["unresolved"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
